"""
Convert manuscript_final.md to a properly formatted Word document.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

SRC = r"C:\Users\mount\.gemini\antigravity\brain\6e0ead93-1a28-4eda-adc8-c75150f49cab\manuscript_final.md"
OUT = r"C:\Users\mount\.gemini\antigravity\scratch\PolicyPrism_Manuscript_Final.docx"

# ── helpers ─────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(para, before=0, after=6, line=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    if line:
        from docx.shared import Pt as P
        fmt.line_spacing = P(line)

def add_heading(doc, text, level):
    """Add a styled heading."""
    sizes   = {1: 16, 2: 14, 3: 12, 4: 11}
    para = doc.add_paragraph()
    para_spacing(para, before=12, after=4)
    run = para.add_run(text)
    set_font(run, size=sizes.get(level, 12), bold=True)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x3A, 0x6B, 0x9F)
    return para

def add_body(doc, text, indent=False):
    """Add a body paragraph, handling **bold** and *italic* inline markup."""
    para = doc.add_paragraph()
    para_spacing(para, before=0, after=6)
    para.paragraph_format.first_line_indent = Inches(0.3) if indent else None
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_inline(para, text)
    return para

def _add_inline(para, text):
    """Parse **bold**, *italic*, and plain text into runs."""
    # Pattern: **bold**, *italic*, plain
    pattern = re.compile(r'(\*\*.*?\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = para.add_run(part)
            set_font(run)

def add_blockquote(doc, text):
    para = doc.add_paragraph()
    para_spacing(para, before=0, after=6)
    para.paragraph_format.left_indent  = Inches(0.5)
    para.paragraph_format.right_indent = Inches(0.5)
    _add_inline(para, text)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para_spacing(para, before=0, after=3)
    para.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    _add_inline(para, text)
    return para

def add_table(doc, headers, rows):
    """Add a nicely formatted table."""
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1A3A5C')
        shading.set(qn('w:color'), 'FFFFFF')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'EBF2FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            _add_inline(cell.paragraphs[0], cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()  # spacer

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def add_separator(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E5C8A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para_spacing(para, before=4, after=4)

# ── Table data ───────────────────────────────────────────────────────────────

TABLE1_HEADERS = ["Concept", "Definition", "Focus", "Example"]
TABLE1_ROWS = [
    ["Strategic Silence",
     "Intentional omission to manage meaning or power\n(Morrison & Milliken, 2000)",
     "Discursive",
     "A regulator omitting enforcement thresholds from public guidance"],
    ["Institutional Voids",
     "Absence of market-supporting institutions\n(Khanna & Palepu, 1997)",
     "Structural-functional",
     "A jurisdiction lacking any data protection authority"],
    ["Ghost Node",
     "Structural absence constitutive of assemblage stability",
     "Ontological-relational",
     "Worker unions absent from the EU AI Act; their foreclosure stabilizes a market/individual-rights logic"],
]

TABLE2_HEADERS = ["Layer", "Lens", "Description"]
TABLE2_ROWS = [
    ["1", "Relationship Extraction", "Actor-to-actor associations using governance verbs with strength scores and provenance flags."],
    ["2", "ANT Tracing", "Classifies intermediaries and mediators; identifies obligatory passage points and translation chains."],
    ["3", "Mediator Classification", "Assesses amplification, blocking, and modification."],
    ["4", "Ecosystem & Ontological Framing", "Assigns edge types (power, logic, ghost) and classifies impacts as constraints or affordances."],
    ["5", "Ontology & Concept Mapping", "Traces power flow, institutional logic, and ghost connections."],
    ["6", "Assemblage Theory", "Evaluates how rules, enforcement, narratives, and resources generate emergent assemblage properties."],
    ["7", "Ghost Node Detection", "Identifies structurally absent actors; operationalizes stratified legibility."],
    ["8", "Comparative Synthesis", "Cross-policy analysis identifying convergent logics, tensions, and divergent patterns."],
]

TABLE3_HEADERS = ["Proposition", "Name", "Description"]
TABLE3_ROWS = [
    ["P1", "Referential Drift", "Portable governance codes carry consistent labels but acquire divergent referents shaped by local institutional logics."],
    ["P2", "Infrastructural Embedding", "Translated codes embedded in operational infrastructures become resistant to reinterpretation in proportion to embedding depth."],
    ["P3", "Stratified Legibility", "Compliance infrastructures produce unequal visibility; counter-translations produce sedimentation or partial deterritorialization."],
]

TABLE4_HEADERS = ["Criterion", "Evidence", "Policy Prism Classification"]
TABLE4_ROWS = [
    ["Functional Relevance",
     "Worker organizations could plausibly evaluate AI system impacts in employment (Annex III, §4), participate in conformity assessments for high-risk workplace AI, and serve on notified-body governance structures.",
     "Satisfied – Governance function clearly exists"],
    ["Textual Invocation",
     "Recital 4 references workers as affected parties; Article 26 requires deployers to inform workers subject to AI decisions; Annex III identifies employment as a high-risk domain.",
     "Satisfied – Invoked as interest-holders; precluded as agents"],
    ["Structural Foreclosure",
     "Conformity assessment (Article 43) requires no worker-organization participation. Notified body designation (Article 33) involves no labor representation. No private right of action for workers in Chapter VII. No worker-organization role in AI Office governance (Article 64–70).",
     "Satisfied – All participation avenues eliminated"],
    ["**Ghost Node Category**", "**Affected-but-Unempowered**", "Confidence: High"],
    ["**Analyst Disagreement Log**",
     "One analyst flagged Article 26 (worker notification) as partial mitigation; consensus position: notification without standing does not constitute governance participation. Disagreement recorded and retained in provenance chain.",
     ""],
]

FIGURE2_HEADERS = ["Actor Category", "EU", "Colorado", "India", "Brazil", "U.S. Federal"]
FIGURE2_ROWS = [
    ["Civil Society / NGOs",
     "Excluded – Rights invoked; NGOs absent",
     "Excluded – Self-testing excludes monitors",
     "Excluded – Framework targets ministries",
     "Silenced – No implementation role",
     "Excluded – Bias/child-protection rhetoric"],
    ["Workers / Labor Unions",
     "Marginalized – No compliance review role",
     "Marginalized – No worker role in decisions",
     "Excluded – No labor-impact analysis",
     "Marginalized – Principle-level only",
     "Marginalized – 'Workers are central' rhetoric"],
    ["End Users",
     "Silenced – Passive aggregate only", "—", "—", "—",
     "Invoked – 'seek redress' only"],
    ["Indigenous Communities",
     "—", "—",
     "Excluded – Universal framing erases differentiated rights",
     "—", "—"],
    ["The Public / Citizens", "—", "—", "—", "—", "Unempowered – 'seek redress' only"],
    ["Researchers and Academics", "—", "—", "—", "—", "Experts/testers, not decision-makers"],
]

# ── Main build ────────────────────────────────────────────────────────────────

import docx

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Title page ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(title, before=60, after=12)
tr = title.add_run("Policy Prism: An Interpretive Artifact for Analyzing\nAI Governance as an Algorithmic Assemblage")
set_font(tr, size=18, bold=True, color=(0x1A, 0x3A, 0x5C))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(subtitle, before=0, after=4)
sr = subtitle.add_run("Manuscript prepared for Information & Organization\nSpecial Issue: Algorithmic Assemblages")
set_font(sr, size=12, italic=True, color=(0x2E, 0x5C, 0x8A))

doc.add_paragraph()
add_separator(doc)
doc.add_paragraph()

# ── Abstract ──────────────────────────────────────────────────────────────────
add_heading(doc, "Abstract", 2)

abstract_paras = [
    "AI governance is frequently treated as a matter of designing rules—specifying what counts as risk, who is responsible, and how compliance is measured. This framing obscures the organizational labor through which governance actually stabilizes. We argue that AI governance is better understood as an *algorithmic assemblage*: a dynamic formation in which legal texts, standards bodies, infrastructure systems, firms, expert groups, and affected populations are coupled through ongoing processes of translation and inscription.",
    "The paper makes two primary contributions. The **Ghost Node** construct identifies actors who are rhetorically invoked and materially affected yet constitutively absent from the compliance architectures through which governance becomes actionable—analytically tied to assemblage stability, not merely absent from its coverage. **Translational Stratification Theory (TST)** is proposed as a middle-range process theory explaining how portable governance vocabularies are institutionalized through three linked stages: referential drift, infrastructural embedding, and stratified legibility. **Policy Prism**, the enabling artifact, uses Large Language Models (LLMs) to generate provisional, evidence-linked analyses of policy documents to support researcher reflection.",
    "Applied to five AI governance frameworks—the EU AI Act, Colorado SB 24-205, India's AI Governance Framework, Brazil's PL 2338/2023, and the U.S. Federal AI Policy Package—the analysis shows that authority concentrates in technocratic nodes, governance is routed through compliance infrastructures, and affected actors are recurrently positioned as Ghost Nodes. Together, Ghost Node and TST provide a theoretical lens and practical instrument for making constitutive absence in algorithmic governance assemblages empirically tractable and contestable.",
]
for p in abstract_paras:
    add_body(doc, p)

add_separator(doc)

# ── Section 1 ─────────────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction", 1)

s1_paras = [
    "Regulatory design for AI can be viewed as essentially a task of making rules: determining types of risk, specifying which party is responsible, and how compliance will be measured. However, as AI enters societal structures, regulatory frameworks evolve through interactions among agencies, standard-setting organizations, infrastructure providers, businesses, professions, and affected populations—not solely through formal written documents.",
    "**The paper's core argument:** AI governance assemblages stabilize through compliance infrastructures before they stabilize in law. As they do, affected actors—civil society, workers, communities subject to algorithmic decisions—are frequently present in the rhetoric of governance but absent from its architecture. Ghost Node and TST provide the analytic tools for studying this pattern comparatively.",
    "We argue that AI governance should be regarded as an Algorithmic Assemblage (Law & Hassard, 1999; Callon et al., 2007): a dynamic formation in which authority, accountability, and meaning emerge and stabilize through ongoing interaction among diverse actors. According to ANT and assemblage theory (DeLanda, 2016), governance is a continuous outcome of translation, inscription, and embedding of categories into technical and organizational systems (Callon, 1986; Latour, 2005). Categories such as \"risk\" and \"fairness\" do not exist as abstractions; they become governance tools only when embodied in standards, audits, and certification processes. Relational theories explain how governance forms, but have been less precise about what is systematically excluded as it spreads—and how that exclusion becomes a condition of durability. Ghost Node and Policy Prism address that gap directly.",
    "We conducted case studies on five AI governance configurations: the EU AI Act; Colorado SB 24-205; India's AI Governance Framework; Brazil's PL 2338/2023 as an exploratory vignette; and the U.S. Federal AI Policy Package. Each features governing nodes at the top of the hierarchy, upward compliance mechanisms, and references to affected populations with no consistent opportunity for durable participatory engagement. From these patterns, we posit TST as a middle-range process theory explaining how portable governance vocabularies redistribute visibility, participation, and authority as they are embedded locally.",
    "The remainder is organized as follows: Section 2 establishes the theoretical rationale for treating AI governance as an assemblage characterized by constitutive absence; Section 3 presents the Policy Prism analytical framework; Section 4 describes the methodology; Section 5 reports results; Section 6 provides the discussion; and Section 7 concludes.",
]
for p in s1_paras:
    add_body(doc, p)

# ── Section 2 ─────────────────────────────────────────────────────────────────
add_heading(doc, "2. AI Governance and the Problem of Absence", 1)

add_body(doc, "ANT directs attention to how diverse actors—policy texts, regulatory agencies, databases, audit routines, standards, model developers, and affected publics—build governance arrangements through translation, inscription, and enrollment (Callon, 1984; Latour, 2005). Assemblage theory extends ANT by emphasizing heterogeneous attractors that solidify systems and lines-of-flight that disrupt them (DeLanda, 2016; Deleuze & Guattari, 1987). Assemblages are constructed through contingent aggregation of material and expressive components; once formed, they remain susceptible to stabilization, disruption, and transformation. AI governance converts concepts such as risk, fairness, and accountability from legal abstractions to governance tools. Once incorporated into standards, registries, certification routes, and enforcement mechanisms, those concepts become durable—but their operationalization differs widely across jurisdictions.")

add_heading(doc, "2.2 The Significance of Structural Absence", 2)

add_body(doc, "AI governance architectures typically acknowledge technological harms while rarely providing affected parties with organizational standing to participate in compliance processes. This is not oversight; structural foreclosure produces specific organizational effects. In Law's (2009) terms, certain actors are \"absent presences\": invoked as reference points but excluded from procedural standing. Floridi's (2013, 2023) account of distributed moral responsibility clarifies how governance architectures render actors structurally invisible while claiming to protect them.")
add_body(doc, "A Ghost Node is identified when an actor's absence is a constitutive condition of assemblage stability—not merely a gap in coverage. Three evidentiary criteria must be simultaneously satisfied:")

add_bullet(doc, "**Functional Relevance:** A plausible governance function exists for which the actor could have been responsible.")
add_bullet(doc, "**Textual Invocation:** The regime identifies the actor as having an interest or standing related to its purposes, but in a way that precludes participation as an agent.")
add_bullet(doc, "**Structural Foreclosure:** The procedural architecture eliminates any avenue for complaint, governance involvement, or enforcement standing.")

doc.add_paragraph()
add_body(doc, "**Table 1. Ghost Node Distinguished from Adjacent Concepts**")
add_table(doc, TABLE1_HEADERS, TABLE1_ROWS)

add_body(doc, "Once all three criteria are met, the node is assigned one of four absence categories: **Affected-but-Unempowered** (mentioned, prohibited from governance); **Rhetorically Invoked** (named in preamble, absent from operative provisions); **Structurally Excluded** (functionally relevant, excluded by actor definition); or **Displaced** (involved in the regulated production stream, outside jurisdictional reach). Policy Prism records analyst agreement and disagreement across all classifications.")

add_heading(doc, "2.3 Policy Prism as Interpretive Artifact", 2)
add_body(doc, "AI governance distributes interpretive power—who determines harm, who provides expertise, and who can contest categories. Policy Prism surfaces structural absences as evidence-linked candidate interpretations for analyst review; they are not findings. The point is to make interpretive disagreement visible as part of the phenomenon, not to eliminate it. Outputs are preliminary, traceable to textual evidence, and designed for human modification.")

add_heading(doc, "2.5 TST: A Middle-Range Process Theory", 2)
add_body(doc, "Drawing on ANT theories of translation (Callon, 1986; Latour, 2005), norm diffusion (Finnemore & Sikkink, 1998), and vernacularization (Merry, 2006), TST describes how portable governance vocabularies—carried by stable labels—acquire divergent referents through local translation, then become durable through organizational embedding. Three stages:")
add_body(doc, "**P1 – Referential Drift:** Labels remain stable as they cross borders; local institutional logics anchor them to divergent referents. The more portable a governance term, the less likely it maintains a stable substantive referent.")
add_body(doc, "**P2 – Infrastructural Embedding:** Drifted vocabularies become durable through incorporation into standards, audits, registries, and certification schemes. Depth of embedding correlates with resistance to reinterpretation.")
add_body(doc, "**P3 – Stratified Legibility:** Compliance infrastructures produce stratified legibility. Actors able to engage with the architecture become visible and actionable; those who cannot are rendered structurally absent. Stage 3 produces either *sedimentation* (challenges absorbed, infrastructure unchanged) or *partial deterritorialization* (genuine redistribution of standing).")
add_body(doc, "TST connects prior work on translation, norm diffusion, and sociomaterial infrastructure (Orlikowski, 2007; Bowker & Star, 1999) through a mechanistic linkage between semantic drift and organizational inequality.")

# ── Section 3 ─────────────────────────────────────────────────────────────────
add_heading(doc, "3. System Overview: The Policy Prism Analytical Framework", 1)
add_body(doc, "Policy Prism provides a multi-layered framework for analyzing governance structures, illustrated in Figure 1. The double-prism architecture derives propositions empirically from decomposition spectra and achieves the reconvergence needed for TST analysis.")
add_body(doc, "*Figure 1. Policy Prism – Analytical Framework for TST of AI Governance.*")
add_body(doc, "**Table 2. Document Analysis: Eight Analytic Layers**")
add_table(doc, TABLE2_HEADERS, TABLE2_ROWS)
add_body(doc, "**Table 3. TST: Testable Propositions**")
add_table(doc, TABLE3_HEADERS, TABLE3_ROWS)

add_heading(doc, "3.2 Technical Implementation", 2)
add_body(doc, "Policy Prism uses multiple LLMs: GPT-5.1 (OpenAI) for governance analysis, assemblage extraction, Ghost Node assessments, and counterfactual tests; GPT-4o-mini for structural parsing and mediator scoring; Google Gemini 1.5 Flash for supplementary web search. No output is treated as a finding; each is an evidence-linked result for researcher reflection. All prompts, model settings, methodological assumptions, and provenance data are retained.")

# ── Section 4 ─────────────────────────────────────────────────────────────────
add_heading(doc, "4. Research Design and Methodology", 1)
add_body(doc, "Policy Prism was developed through an echeloned approach to e-Design Science Research (eDSR) (Tuunanen et al., 2024) across four stages: conceptual framing; first version; second version; and comparative application to five AI governance configurations. An interpretivist paradigm was adopted: the goal is to demonstrate that Policy Prism renders governance relations, accountability paths, and structural absences visible and amenable to challenge. Policy documents are treated as sites where governance logics, exclusions, and accountability relationships are textually organized—not as transparent expressions of institutional intent.")

add_heading(doc, "4.1 Methodology Validation", 2)
add_body(doc, "Validation proceeds across three dimensions: **Traceability** (analytic claims reference specific textual evidence; model outputs are revisable provisional interpretations, not definitive results); **Contestability** (outputs are flagged with uncertainty metrics; analysts can override or annotate; disagreements are documented); and **Theoretical Coherence** (design decisions constrained by ANT and assemblage theory; premature abstractions removed). This study demonstrates Policy Prism interpretively across five cases; inter-rater reliability studies at scale, external evaluator comparisons, and longitudinal deployment remain future work.")

add_heading(doc, "4.2 Data Sources", 2)
add_body(doc, "This research examines four primary configurations—the EU AI Act, Colorado SB 24-205, India's AI Governance Guidelines, and the U.S. Federal AI Policy Package—plus Brazil's PL 2338/2023 as an exploratory vignette. The U.S. materials (America's AI Action Plan, July 2025; National Policy Framework for AI: Legislative Recommendations, March 2026) are treated as an emergent federal assemblage, not a settled statutory regime. Brazil-derived claims are provisional, as the text was translated via secondary summaries. A reflexive assessment panel was completed before each Ghost Node analysis, with analyst disagreements recorded.")

# ── Section 5 ─────────────────────────────────────────────────────────────────
add_heading(doc, "5. Results", 1)

add_heading(doc, "5.1 Ghost Node Analysis", 2)
add_body(doc, "Figure 2 presents Ghost Node findings across all five jurisdictions. Civil society, workers, and affected communities are structurally absent in every case despite being rhetorically invoked. The pattern confirms that infrastructural embedding creates systematic inequalities in visibility, participation, and accountability.")
add_body(doc, "**Figure 2. Ghost Node Exclusion Matrix: Structural absences across five AI governance jurisdictions**")
add_table(doc, FIGURE2_HEADERS, FIGURE2_ROWS)

add_body(doc, "To illustrate methodological concreteness, Table 4 presents the complete evidentiary classification for one Ghost Node: worker unions under the EU AI Act.")
add_body(doc, "**Table 4. Illustrative Ghost Node Classification: Worker Unions, EU AI Act**")
add_table(doc, TABLE4_HEADERS, TABLE4_ROWS)

add_body(doc, "Policy Prism's provenance chain for this classification retains: the verbatim article citations above; the prompt used to generate the initial candidate interpretation; the model's uncertainty flag on Article 26; the analyst override annotation; and the final evidence-graded classification. This chain is fully inspectable and revisable at each stage.")

add_heading(doc, "5.2 The U.S. Case", 2)
add_body(doc, "Because the source documents represent proposals for legislative action rather than a settled statutory regime, the U.S. is treated as an emergent federal assemblage requiring interpretive caution.")
add_body(doc, "Power is constructed through inscriptions—executive orders, legislative recommendations, contracting mechanisms, standards development, and agency coordination—that produce an authoritative center without a single regulatory node. The Action Plan calls for conditions enabling \"private sector-led innovation\" while asserting U.S. leadership as \"the gold standard for AI worldwide\"; the 2026 Framework recommends against creating \"any new Federal rule-making body,\" relying instead on \"industry-led standards.\" Federalization is accomplished through stratified authority across executive, technical, and industrial nodes, with preemptive language consolidating \"one national standard\" to prevent state law \"from impeding the implementation of the U.S. Strategy for AI Global Leadership.\"")
add_body(doc, "The Ghost Node pattern follows. Workers are rhetorically centered yet expected to adapt through skills and training rather than participate in compliance design or governance. The public may pursue \"redress\" but holds no standing in procurement, standard-setting, or model evaluation. Researchers and academics function as knowledge sources rather than decision participants. Non-U.S. populations are expected to adopt U.S. standards and platforms.")
add_body(doc, "All three TST propositions are supported. Risk, Innovation, Evaluation, Science, and Security are framed as neutral yet embedded in a security-industrial frame through procurement structures, standards frameworks, and export controls. The primary exclusion mechanism is selective inclusion: actors are mentioned but lack standing, enrolled but lack authority.")

add_heading(doc, "5.3 Cross-Policy Themes", 2)
add_body(doc, "Each regime defines its purposes in public-good terms while exercising authority through infrastructures that direct decision-making upward. The organizational result is consistent across all five: entities controlling obligatory passage points gain standing; entities affected by governance rarely receive commensurate power.")
add_body(doc, "**Accountability Architecture** (Figure 3). Accountability flows upward—Deployer to Regulator, Operator to Certification Body, Provider to Standards Bodies—rather than toward Affected Communities. This exemplifies what Shaikh et al. (2023) call \"relational spaces\" of algorithmic accountability: compliance architectures route accountability rather than generating enforceable obligations to those experiencing AI governance's effects.")
add_body(doc, "**Authority Concentration** (Figure 4). Apex authority resides in technocratic bodies in all five jurisdictions: the European Commission and AI Office (CEN/CENELEC standards); Colorado's Attorney General (sole enforcement, no private right of action); India's CARO, AIREC, and IndiaAI Mission; Brazil's ANPD; and the U.S.'s OMB, CAISI, and CAIOC. In each case, authoritative interpretation resides at the top of the assemblage.")
add_body(doc, "**Worker Exclusion** (Figure 5). Worker organizations are structurally foreclosed across all five regimes despite each governing domains where AI affects employment and workplace decision-making. The U.S. case is sharpest: a \"worker-first AI agenda\" coexists with no institutional role for labor organizations in compliance design, evaluation, or governance architecture.")
add_body(doc, "**Governance Styles** (Figure 6). The five regimes organize and legitimate governance differently—EU through market convergence; Colorado through consumer protection and AG enforcement; India through developmental coordination; Brazil through civic-developmental language; the U.S. through security-industrial competition. Despite these differences, accountability flows through compliance architectures rather than participatory co-governance in all five cases.")

add_heading(doc, "5.4 TST Applied: Evidence Across Five Cases", 2)
add_body(doc, "Applying the three-stage pipeline described in Section 2.5, the comparative data produce the following findings.")
add_body(doc, "**P1 – Referential Drift.** \"Risk\" names different objects across jurisdictions: product safety (EU), social harm (Brazil), developmental coordination (India), consequential consumer decisions (Colorado), and security-industrial competition (U.S.). Shared vocabulary conceals institutional divergence; apparent terminological convergence leaves structural inequality intact.")
add_body(doc, "**P2 – Infrastructural Embedding.** In each case, translated risk categories become durable through apex institutions acting as obligatory passage points—the AI Office and CEN/CENELEC (EU), the Attorney General (Colorado), CARO and IndiaAI Mission (India), ANPD (Brazil), and OMB/CAISI/CAIOC (U.S.). Actors able to navigate these structures become visible; those who cannot are marginalized.")
add_body(doc, "**P3 – Stratified Legibility.** Ghost Nodes recur across all five cases. Counter-translational capacity exists unevenly: Brazil's ANPD—with formal reclassification authority—is the clearest case of institutional capacity for partial deterritorialization; elsewhere, challenges are largely absorbed through sedimentation.")
add_body(doc, "**Convergence Finding.** Shared vocabulary consistently yields the same organizational effects: authority concentrating in technocratic nodes, governance routed through compliance infrastructures, and affected actors structurally absent. TST is therefore a recurrent organizational feature of algorithmic governance assemblages, not an anomaly of particular jurisdictions.")
add_body(doc, "**Worker Vignette** (Figure 7). The worker case (detailed in Table 4 above) illustrates the pipeline: labor concerns are displaced into other policy idioms (P1); compliance pathways provide no durable role for worker organizations (P2); workers remain visible as policy objects but hold no participatory standing (P3). The classification and its provenance chain are published in full in Appendix B.")

# ── Section 6 ─────────────────────────────────────────────────────────────────
add_heading(doc, "6. Discussion", 1)

add_heading(doc, "6.1 Translational Stratification Theory", 2)
add_body(doc, "Algorithmic governance stabilizes through organizations before it stabilizes in law—through standards bodies that define conformity before legislatures codify it, procurement offices that embed risk categories before agencies formalize them, and audit regimes that convert abstractions into auditable routines. TST (Section 2.5) specifies this as an empirically traceable pipeline from referential drift through infrastructural embedding to stratified legibility.")
add_body(doc, "TST's theoretical contribution relative to prior work on translation (Callon, 1986; Latour, 2005), norm diffusion (Finnemore & Sikkink, 1998), and sociomaterial infrastructure (Orlikowski, 2007; Bowker & Star, 1999) is to specify *what is structurally foreclosed* as diffusion proceeds, and how that foreclosure becomes constitutive of durability rather than incidental to it. The U.S. case extends this: infrastructure—not centralization—is the operative mechanism of stratification. TST also reframes ethics washing: terms like \"risk\" and \"fairness\" persist not merely through rhetorical flexibility, but because compliance routines convert that flexibility into durable sociotechnical reality resistant to external contestation (Floridi, 2019, 2023).")
add_body(doc, "Ghost Node analysis (Section 2.2) renders the organizational costs of that hardening visible by identifying constitutive absences—structural positions analytically tied to assemblage stability—rather than mapping stakeholder gaps.")

add_heading(doc, "6.2 Conceptual Contributions", 2)
add_body(doc, "**Referential drift** establishes that vocabulary-level harmonization leaves structural divergence intact. Shared labels do not produce shared governance objects; harmonization operating only terminologically cannot address the inequalities drift produces.")
add_body(doc, "**Infrastructural closure** identifies where translated categories become resistant to contestation. Closure is most complete in the EU's CE-marking and notified-body regime; Brazil's ANPD demonstrates it is a variable, not an inevitability.")
add_body(doc, "**Ghost Node** shifts the analytic question from \"who is absent?\" to \"how does this absence organize the assemblage?\" Reform must engage the access criteria, evidentiary thresholds, and standing requirements through which structural absence is produced—not merely supplement them with participation add-ons. Three organizational implications follow: governance vocabularies constitute organizational fields through standards and procurement rather than legislative fiat; compliance infrastructures concentrate authority by engineering obligatory passage points; and stabilization and structural absence are co-constitutive.")

add_heading(doc, "6.3 Contributions to the Information & Organization Conversation", 2)
add_body(doc, "For *Information & Organization* scholarship on algorithmic assemblages, the central contribution is to show that assemblages stabilize through the ongoing coupling of portable vocabularies to material-organizational infrastructures—not through discourse alone. Standards, audits, procurement systems, registries, and coordinating bodies determine who counts, who can intervene, and which interpretations become authoritative (Orlikowski, 2007; DeLanda, 2016).")
add_body(doc, "**Ghost Node is the primary conceptual contribution.** It is not simply missing representation; it is absence that is analytically tied to the stability of the assemblage—structural foreclosure that sustains the governance formation rather than reflecting gaps in coverage. This distinguishes it from strategic silence, institutional voids, and stakeholder omission. **TST is the supporting process theory**, specifying the sequenced mechanism (referential drift → infrastructural embedding → stratified legibility) that makes the finding generalizably comparative. **Policy Prism is the enabling artifact**, operationalizing both through traceable, contestable, evidence-linked analysis.")
add_body(doc, "Policy Prism is demonstrated interpretively across five cases here; broader validation—inter-rater reliability at scale, external evaluator comparisons, longitudinal deployment—remains future work. Three design principles proved essential: *traceability* (classifications link to textual evidence), *contestability* (interpretations remain open to challenge), and *productive friction* (ambiguity preserved until evidence warrants resolution). The LLM functions as an analytical accelerator, not a decision-maker, with disciplined prompting and provenance chains embedding reflexivity as an infrastructural property (Bender et al., 2021; Zheng & Lee, 2023).")
add_body(doc, "Three LLM-specific risks require acknowledgment: *probabilistic closure* (filling structural absences with plausible inferences); *hegemonic training corpora* (Global North regulatory texts dominate, reducing salience of non-institutional actors); and *semantic compression* (rhetorical mention and procedural standing may be collapsed). Policy Prism addresses these through evidence-only output requirements, missing-signal diagnostics, fragility scoring, and provenance tracking. Human review and structured contestation are constitutive of the artifact's epistemic validity (Paseri & Durante, 2025).")
add_body(doc, "**Policy Prism as an information artifact.** A distinctive contribution to I&O's concern with information—not only organization—is that Policy Prism makes governance *legibility* itself an object of analysis. Conventional compliance reporting generates upward-flowing information: risk classifications, conformity assessments, and audit results that confirm the system is working for those at its apex. Policy Prism redirects analytical attention to the information that does not flow: the absence signals, foreclosed standing pathways, and missing-signal diagnostics that reveal whom governance has rendered structurally invisible. In this sense, it functions as a counter-infrastructure of information—one designed not to confirm compliance but to surface the organizational conditions under which certain actors cannot produce legible claims at all. This connects directly to I&O's sustained interest in the politics of visibility in information systems (Introna & Nissenbaum, 2000), the classificatory labor that constitutes organizational fields (Bowker & Star, 1999), and the sociomaterial production of organizational sensemaking (Weick, 1995; Orlikowski & Scott, 2008).")
add_body(doc, "This framing also opens a practical research agenda. Recent I&O scholarship on algorithmic management, platform governance, and data infrastructures has documented how information asymmetries compound organizational inequality (Zuboff, 2019; Kallinikos et al., 2013). Ghost Node analysis extends that line by showing that informational invisibility is not merely a byproduct of governance design—it is a constitutive mechanism through which assemblage stability is achieved and maintained. Policy Prism provides a methodologically explicit way to study this mechanism comparatively, making it a resource for I&O researchers working across regulatory regimes, organizational contexts, and levels of analysis.")

# ── Section 7 ─────────────────────────────────────────────────────────────────
add_heading(doc, "7. Conclusion", 1)
add_body(doc, "Ghost Node and TST together show that algorithmic governance assemblages achieve durability through organizational arrangements that translate portable vocabularies into compliance infrastructures while rendering selected actors differentially visible, legible, and contestable. Shared vocabulary travels easily; local translation concentrates authority in technocratic nodes, routes governance through compliance architectures, and positions civil society, workers, and affected communities as Ghost Nodes whose structural absence sustains assemblage stability. Counter-translation pathways exist—Brazil's ANPD is the clearest example—but the capacity to use them is produced by the very arrangements that generate structural foreclosure.")
add_body(doc, "Assemblage durability arises not only from what is formalized, but from what is systematically deferred, displaced, or structurally foreclosed. Addressing those foreclosures—rather than supplementing them—is the design challenge that more equitable algorithmic governance requires.")

add_heading(doc, "Limitations", 3)
add_body(doc, "This study demonstrates TST and Ghost Node analysis interpretively across five cases. Brazil-derived claims are provisional. The U.S. case draws on proposal-based documents representing an emergent federal assemblage. Policy Prism's validation is theory-building and interpretive; inter-rater reliability studies at scale and external comparisons remain future work.")

add_separator(doc)

# ── References ────────────────────────────────────────────────────────────────
add_heading(doc, "References", 1)

refs = [
    "Bender, E. M., Gebru, T., McMillan-Major, A., & Shmitchell, S. (2021). On the dangers of stochastic parrots: Can language models be too big? *FAccT '21* (pp. 610–623). ACM.",
    "Bowker, G. C., & Star, S. L. (1999). *Sorting things out: Classification and its consequences.* MIT Press.",
    "Callon, M. (1984). Some elements of a sociology of translation. *The Sociological Review, 32*(1_suppl), 196–233.",
    "Callon, M. (1986). The sociology of an actor-network: The case of the electric vehicle. In M. Callon, J. Law, & A. Rip (Eds.), *Mapping the dynamics of science and technology* (pp. 19–34). Macmillan.",
    "DeLanda, M. (2016). *Assemblage theory.* Edinburgh University Press.",
    "Deleuze, G., & Guattari, F. (1987). *A thousand plateaus: Capitalism and schizophrenia* (B. Massumi, Trans.). University of Minnesota Press.",
    "Finnemore, M., & Sikkink, K. (1998). International norm dynamics and political change. *International Organization, 52*(4), 887–917.",
    "Floridi, L. (2013). *The ethics of information.* Oxford University Press.",
    "Floridi, L. (2019). Translating principles into practices of digital ethics. *Philosophy & Technology, 32*(2), 185–193.",
    "Floridi, L. (2023). *The ethics of artificial intelligence.* Oxford University Press.",
    "Introna, L. D., & Nissenbaum, H. (2000). Shaping the web: Why the politics of search engines matters. *The Information Society, 16*(3), 169–185.",
    "Kallinikos, J., Aaltonen, A., & Marton, A. (2013). The ambivalent ontology of digital artifacts. *MIS Quarterly, 37*(2), 357–370.",
    "Khanna, T., & Palepu, K. G. (1997). Why focused strategies may be wrong for emerging markets. *Harvard Business Review, 75*(4), 41–48.",
    "Lampland, M., & Star, S. L. (Eds.) (2009). *Standards and their stories: How quantifying, classifying, and formalizing practices shape everyday life.* Cornell University Press.",
    "Latour, B. (2005). *Reassembling the social: An introduction to actor-network-theory.* Oxford University Press.",
    "Law, J. (2009). Actor network theory and material semiotics. In B. S. Turner (Ed.), *The new Blackwell companion to social theory* (pp. 141–158). Wiley-Blackwell.",
    "Merry, S. E. (2006). *Human rights and gender violence: Translating international law into local justice.* University of Chicago Press.",
    "Morrison, E. W., & Milliken, F. J. (2000). Organizational silence. *Academy of Management Review, 25*(4), 706–725.",
    "Orlikowski, W. J. (2007). Sociomaterial practices. *Organization Studies, 28*(9), 1435–1448.",
    "Orlikowski, W. J., & Scott, S. V. (2008). Sociomateriality: Challenging the separation of technology, work and organization. *Academy of Management Annals, 2*(1), 433–474.",
    "Paseri, L., & Durante, M. (2025). [Reference for LLM epistemological risks in governance contexts].",
    "Peck, J., & Theodore, N. (2010). Mobilizing policy: Models, methods, and mutations. *Geoforum, 41*(2), 169–174.",
    "Shaikh, M., et al. (2023). Relational spaces of algorithmic accountability. *Information and Organization.*",
    "Stone, D. (2012). Transfer and translation of policy. *Policy Studies, 33*(6), 483–499.",
    "Timmermans, S., & Epstein, S. (2010). A world of standards but not a standard world. *Annual Review of Sociology, 36*, 69–89.",
    "Tuunanen, T., Winter, R., & vom Brocke, J. (2024). Dealing with complexity in design science research. *MIS Quarterly, 48*(2), 427–458.",
    "Weick, K. E. (1995). *Sensemaking in organizations.* Sage.",
    "The White House. (2025a). *America's AI Action Plan.* The White House.",
    "The White House. (2026). *National policy framework for artificial intelligence: Legislative recommendations.* The White House.",
    "Zheng, Y., & Lee, A. S. (2023). [Reference for LLM reflexivity and auditability in research practice].",
    "Zuboff, S. (2019). *The age of surveillance capitalism.* PublicAffairs.",
]

for ref in refs:
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=3)
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    _add_inline(p, ref)

doc.save(OUT)
print(f"Saved: {OUT}")
