"""
Generate TST Three-Stage Pipeline as a Word table.
Core claim: Labels travel, referents drift, infrastructures stabilize inequality.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    c = color_hex.lstrip('#')
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), c)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            el.set(qn('w:space'), '0')
            borders.append(el)
    tcPr.append(borders)

def add_run(paragraph, text, bold=False, italic=False, size=10, color=None, font='Calibri'):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        c = color.lstrip('#')
        run.font.color.rgb = RGBColor(int(c[0:2],16), int(c[2:4],16), int(c[4:6],16))
    return run

def set_cell_vertical_alignment(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), align)
    tcPr.append(va)

def set_cell_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width.twips)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# ──────────────────────────────────────────────────────────────
# DATA
# ──────────────────────────────────────────────────────────────

STAGES = [
    {
        'num': '1', 'name': 'Referential Drift', 'prop': 'P1',
        'mechanism': (
            'Portable governance terms ("risk," "accountability," "high-risk AI") '
            'travel globally but their substantive referents shift in each local context. '
            'Semantic convergence masks substantive divergence.'
        ),
        'evidence': {
            'EU AI Act': '"Risk" = product safety (Annex III enumeration)',
            'Brazil PL 2338': '"Risk" = adaptive social harm (objective liability)',
            'India NITI Aayog': '"Risk" = ministerial priorities (developmental resilience)',
            'Colorado SB 205': '"Risk" = consequential decisions (consumer harm via AG)',
        },
        'prop_text': 'Labels travel, referents shift across jurisdictions',
        'bg': 'EFF6FF', 'accent': '2563EB',
    },
    {
        'num': '2', 'name': 'Infrastructural Embedding', 'prop': 'P2',
        'mechanism': (
            'Drifted categories are locked into compliance infrastructures — '
            'conformity assessments, standards, registries, and administrative routines. '
            'Apex nodes gain interpretive authority through infrastructural position.'
        ),
        'evidence': {
            'EU AI Act': 'CE marking, notified bodies, EU database — strong infrastructural closure',
            'Brazil PL 2338': 'ANPD oversight, LGPD-derived audit routines',
            'India NITI Aayog': 'Principles-based only; no binding compliance infrastructure',
            'Colorado SB 205': 'AG rulemaking, duty of care; no private right of action',
        },
        'prop_text': 'Drifted categories harden into auditable routines that resist revision',
        'bg': 'ECFDF5', 'accent': '059669',
    },
    {
        'num': '3', 'name': 'Stratified Legibility', 'prop': 'P3',
        'mechanism': (
            'Compliance infrastructures unevenly distribute visibility. '
            'Regulators, auditors, and providers become legible; workers, communities, '
            'and end users become ghost nodes — structurally peripheral.'
        ),
        'evidence': {
            'EU AI Act': 'Regulators & auditors visible; end users invisible (CE system)',
            'Brazil PL 2338': 'ANPD provides formal counter-translation channel (unique)',
            'India NITI Aayog': 'Informal pathways only; no institutional channel for affected populations',
            'Colorado SB 205': 'AG sole gatekeeper; no civil society standing',
        },
        'prop_text': 'Infrastructures stabilize inequality of visibility and power',
        'bg': 'FFFBEB', 'accent': 'D97706',
    },
]

JURISDICTIONS = ['EU AI Act', 'Brazil PL 2338', 'India NITI Aayog', 'Colorado SB 205']

# ──────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ──────────────────────────────────────────────────────────────

doc = Document()

# Page setup — landscape
for section in doc.sections:
    section.orientation = 1  # landscape
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

# ─── Title ───
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(title, 'Table X. ', bold=True, size=11, color='#374151')
add_run(title, 'Translational Stratification Theory — Three-Stage Analytical Pipeline', bold=True, size=11, color='#111827')
title.paragraph_format.space_after = Pt(2)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(subtitle, 'Core claim: Labels travel, referents drift, infrastructures stabilize inequality.', italic=True, size=10, color='#6B7280')
subtitle.paragraph_format.space_after = Pt(8)

# ─── Main Table ───
# Columns: Stage | Mechanism | EU | Brazil | India | Colorado | Proposition
num_cols = 7
table = doc.add_table(rows=1, cols=num_cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Set table style to remove default borders, we'll add our own
table.style = 'Table Grid'

# ─── Header Row ───
hdr = table.rows[0]
hdr_texts = ['Stage', 'Mechanism', 'EU AI Act', 'Brazil PL 2338', 'India NITI Aayog', 'Colorado SB 205', 'Proposition']
hdr_widths = [Inches(1.1), Inches(2.4), Inches(1.6), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.2)]

for ci, (txt, w) in enumerate(zip(hdr_texts, hdr_widths)):
    cell = hdr.cells[ci]
    set_cell_shading(cell, '#1E293B')
    set_cell_vertical_alignment(cell, 'center')
    set_cell_width(cell, w)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, txt, bold=True, size=9, color='#F8FAFC')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# Set row height for header
hdr.height = Pt(28)

# ─── Data Rows ───
for stage in STAGES:
    row = table.add_row()

    # Stage cell
    c0 = row.cells[0]
    set_cell_shading(c0, stage['bg'])
    set_cell_vertical_alignment(c0, 'center')
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, stage['num'], bold=True, size=18, color=f'#{stage["accent"]}')
    p0b = c0.add_paragraph()
    p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0b, stage['name'], bold=True, size=9, color=f'#{stage["accent"]}')
    p0b.paragraph_format.space_before = Pt(2)

    # Mechanism cell
    c1 = row.cells[1]
    set_cell_vertical_alignment(c1, 'center')
    p1 = c1.paragraphs[0]
    add_run(p1, stage['mechanism'], size=9, color='#334155')
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)

    # Jurisdiction cells (4 columns)
    for ji, jname in enumerate(JURISDICTIONS):
        cell = row.cells[2 + ji]
        set_cell_vertical_alignment(cell, 'center')
        p = cell.paragraphs[0]
        evidence = stage['evidence'][jname]
        add_run(p, evidence, size=9, color='#334155')
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # Proposition cell
    c6 = row.cells[6]
    set_cell_shading(c6, 'EEF2FF')
    set_cell_vertical_alignment(c6, 'center')
    p6 = c6.paragraphs[0]
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p6, stage['prop'], bold=True, size=14, color='#4F46E5')
    p6b = c6.add_paragraph()
    p6b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p6b, stage['prop_text'], italic=True, size=8, color='#4338CA')
    p6b.paragraph_format.space_before = Pt(2)

# ─── Outcomes Row (merged across all columns) ───
out_row = table.add_row()

# Merge first 3 cells for Sedimentation
sedi_cell = out_row.cells[0]
sedi_cell.merge(out_row.cells[3])
set_cell_shading(sedi_cell, 'EEF2FF')
p_s = sedi_cell.paragraphs[0]
add_run(p_s, 'Sedimentation ', bold=True, size=10, color='#4338CA')
add_run(p_s, '(feeds back to Stage 2)', italic=True, size=9, color='#4338CA')
p_s2 = sedi_cell.add_paragraph()
add_run(p_s2, 'Categories stabilize across systems → infrastructural closure → resistance shifts to recoding rather than categorical rejection.', size=9, color='#4338CA')
p_s.paragraph_format.space_before = Pt(4)
p_s2.paragraph_format.space_after = Pt(4)

# Merge remaining cells for Counter-Translation
ct_cell = out_row.cells[4]
ct_cell.merge(out_row.cells[6])
set_cell_shading(ct_cell, 'FEF2F2')
p_c = ct_cell.paragraphs[0]
add_run(p_c, 'Counter-Translation ', bold=True, size=10, color='#991B1B')
add_run(p_c, '(feeds back to Stage 1)', italic=True, size=9, color='#991B1B')
p_c2 = ct_cell.add_paragraph()
add_run(p_c2, 'Ghost nodes reframe governance terms — redefining "risk" as collective harm rather than manageable product defect. Reopens referential drift.', size=9, color='#B91C1C')
p_c.paragraph_format.space_before = Pt(4)
p_c2.paragraph_format.space_after = Pt(4)

# ─── Note below table ───
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(8)
add_run(note, 'Note. ', bold=True, italic=True, size=9, color='#6B7280')
add_run(note, 'Stage 3 produces two possible outcomes: sedimentation (categories stabilize, feeding back to Stage 2) or counter-translation (ghost nodes contest dominant codings, feeding back to Stage 1). P1–P3 correspond to the three formal propositions of Translational Stratification Theory.', italic=True, size=9, color='#6B7280')

# ──────────────────────────────────────────────────────────────
# SAVE
# ──────────────────────────────────────────────────────────────

output = r'c:\Users\mount\.gemini\antigravity\scratch\TST_Table.docx'
doc.save(output)
print(f'[OK] Saved: {output}')
